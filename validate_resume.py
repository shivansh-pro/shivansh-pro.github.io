#!/usr/bin/env python3
"""
Resume ATS Validation Script

Validates that generated PDF and DOCX resume files are:
- Openable and non-corrupt
- Text-extractable (not image-based)
- Free of em dashes and en dashes
- Structurally sound for ATS parsing
- Contain expected sections and keywords

Run before committing: python validate_resume.py
"""

import sys
from pathlib import Path

try:
    import fitz  # PyMuPDF
except ImportError:
    print("FAIL: PyMuPDF not installed. Run: pip install PyMuPDF")
    sys.exit(1)

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("FAIL: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


SITE_DIR = Path(__file__).parent / "_site"
PDF_PATH = SITE_DIR / "Resume-ShivanshSingh.pdf"
DOCX_PATH = SITE_DIR / "Resume-ShivanshSingh.docx"

# Keywords that must appear in extracted text (case-insensitive)
REQUIRED_KEYWORDS = [
    "shivansh singh",
    "pro.shivansh@gmail.com",
    "data engineer",
    "professional experience",
    "education",
]

# Section headers ATS systems look for
EXPECTED_SECTIONS = [
    "professional summary",
    "professional experience",
    "education",
    "certifications",
    "side projects",
]

# Characters that can break ATS parsing
# Smart quotes are fine for modern ATS; em/en dashes are the real risk
FORBIDDEN_CHARS = {
    "\u2014": "em dash",
    "\u2013": "en dash",
}


def check_file_exists(path: Path, label: str) -> bool:
    if not path.exists():
        print(f"FAIL [{label}]: File not found: {path}")
        return False
    size = path.stat().st_size
    if size == 0:
        print(f"FAIL [{label}]: File is empty (0 bytes)")
        return False
    if size > 10_000_000:
        print(f"WARN [{label}]: File is unusually large ({size / 1_000_000:.1f} MB)")
    return True


def validate_pdf() -> list[str]:
    errors = []
    label = "PDF"

    if not check_file_exists(PDF_PATH, label):
        return [f"{label}: File missing or empty"]

    try:
        doc = fitz.open(str(PDF_PATH))
    except Exception as e:
        return [f"{label}: Cannot open file: {e}"]

    # Page count
    if doc.page_count == 0:
        errors.append(f"{label}: No pages found")
        doc.close()
        return errors

    print(f"  {label}: {doc.page_count} page(s), {PDF_PATH.stat().st_size / 1024:.0f} KB")

    # Extract all text
    full_text = ""
    for page_num in range(doc.page_count):
        page = doc[page_num]
        page_text = page.get_text()
        if not page_text.strip():
            errors.append(f"{label}: Page {page_num + 1} has no extractable text (may be image-based)")
        full_text += page_text + "\n"

    doc.close()

    if not full_text.strip():
        errors.append(f"{label}: No text extractable from entire document")
        return errors

    text_lower = full_text.lower()

    # Check required keywords
    for keyword in REQUIRED_KEYWORDS:
        if keyword.lower() not in text_lower:
            errors.append(f"{label}: Missing required keyword: '{keyword}'")

    # Check expected sections
    for section in EXPECTED_SECTIONS:
        if section.lower() not in text_lower:
            errors.append(f"{label}: Missing expected section: '{section}'")

    # Check forbidden characters
    for char, name in FORBIDDEN_CHARS.items():
        if char in full_text:
            count = full_text.count(char)
            errors.append(f"{label}: Found {count} {name}(s) - these can break ATS parsing")

    return errors


def validate_docx() -> list[str]:
    errors = []
    label = "DOCX"

    if not check_file_exists(DOCX_PATH, label):
        return [f"{label}: File missing or empty"]

    try:
        doc = Document(str(DOCX_PATH))
    except Exception as e:
        return [f"{label}: Cannot open file: {e}"]

    # Extract all text from paragraphs
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += cell.text + "\n"

    print(f"  {label}: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} table(s), {DOCX_PATH.stat().st_size / 1024:.0f} KB")

    if not full_text.strip():
        errors.append(f"{label}: No text extractable from document")
        return errors

    text_lower = full_text.lower()

    # Check required keywords
    for keyword in REQUIRED_KEYWORDS:
        if keyword.lower() not in text_lower:
            errors.append(f"{label}: Missing required keyword: '{keyword}'")

    # Check expected sections
    for section in EXPECTED_SECTIONS:
        if section.lower() not in text_lower:
            errors.append(f"{label}: Missing expected section: '{section}'")

    # Check forbidden characters
    for char, name in FORBIDDEN_CHARS.items():
        if char in full_text:
            count = full_text.count(char)
            errors.append(f"{label}: Found {count} {name}(s) - these can break ATS parsing")

    # Check heading styles (ATS uses these for section detection)
    heading_styles_used = set()
    for para in doc.paragraphs:
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            heading_styles_used.add(para.style.name)

    if not heading_styles_used:
        errors.append(f"{label}: No Heading styles found. ATS relies on Word heading styles for section detection.")
    else:
        print(f"  {label}: Heading styles found: {', '.join(sorted(heading_styles_used))}")

    return errors


def validate_source() -> list[str]:
    """Check the source .qmd file for forbidden characters."""
    errors = []
    label = "SOURCE"
    qmd_path = Path(__file__).parent / "resume.qmd"

    if not qmd_path.exists():
        return [f"{label}: resume.qmd not found"]

    content = qmd_path.read_text(encoding="utf-8")

    # Check for -- (which Quarto renders as em dash)
    # Exclude YAML frontmatter and code blocks
    in_frontmatter = False
    in_code_block = False
    for line_num, line in enumerate(content.split("\n"), 1):
        if line.strip() == "---":
            in_frontmatter = not in_frontmatter
            continue
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_frontmatter or in_code_block:
            continue
        if " -- " in line:
            errors.append(f"{label}: Line {line_num}: Found ' -- ' (renders as em dash): ...{line.strip()[:80]}...")

    # Check for Unicode em/en dashes in source
    for char, name in FORBIDDEN_CHARS.items():
        if char in content:
            for line_num, line in enumerate(content.split("\n"), 1):
                if char in line:
                    errors.append(f"{label}: Line {line_num}: Found {name} in source")

    return errors


def main():
    print("=" * 60)
    print("Resume ATS Validation")
    print("=" * 60)

    all_errors = []

    print("\n[1/3] Checking source file (resume.qmd)...")
    source_errors = validate_source()
    all_errors.extend(source_errors)
    if not source_errors:
        print("  SOURCE: All checks passed")

    print(f"\n[2/3] Checking PDF ({PDF_PATH.name})...")
    pdf_errors = validate_pdf()
    all_errors.extend(pdf_errors)
    if not pdf_errors:
        print("  PDF: All checks passed")

    print(f"\n[3/3] Checking DOCX ({DOCX_PATH.name})...")
    docx_errors = validate_docx()
    all_errors.extend(docx_errors)
    if not docx_errors:
        print("  DOCX: All checks passed")

    print("\n" + "=" * 60)
    if all_errors:
        print(f"VALIDATION FAILED: {len(all_errors)} issue(s) found\n")
        for error in all_errors:
            print(f"  - {error}")
        print()
        sys.exit(1)
    else:
        print("VALIDATION PASSED: All checks passed")
        print()
        sys.exit(0)


if __name__ == "__main__":
    main()
